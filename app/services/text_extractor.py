"""Text extraction from PDF and DOCX files.

Text extraction module.
Supports pdfplumber, PyMuPDF, and python-docx backends.
"""

import io
import logging
from pathlib import Path
from typing import Optional

import pdfplumber
import fitz  # PyMuPDF
from docx import Document

logger = logging.getLogger(__name__)


class TextExtractionError(Exception):
    """Raised when text extraction fails."""


def _clean_text(text: str) -> str:
    """Remove invisible Unicode characters that garble extracted text.

    PDF extractors (especially pdfplumber) sometimes insert zero-width spaces
    (\u200b), zero-width joiners (\u200d), and other invisible Unicode
    characters between visible characters. These break sentence splitting,
    citation regex, and LLM text comprehension.

    Also collapses excessive whitespace from the removal.

    PDF line-wrap reconstruction: PDF text extraction emits a newline at every
    visual line wrap, so a single paragraph becomes many "paragraphs". Real
    paragraph breaks are the double-newlines the extractors already emit (and
    that DOCX extraction produces natively). We join single newlines into
    spaces while preserving ``\\n\\n`` paragraph boundaries, so downstream
    paragraph splitters see real paragraphs. This is a no-op for DOCX (which
    has no single-newline line-wraps). Hyphenated line-breaks ("represen-\\n
    tation") are de-hyphenated; other single-newlines become a single space.
    """
    import re

    # Remove zero-width characters
    text = text.replace("\u200b", "")  # zero-width space
    text = text.replace("\u200c", "")  # zero-width non-joiner
    text = text.replace("\u200d", "")  # zero-width joiner
    text = text.replace("\ufeff", "")  # byte order mark / zero-width no-break space
    text = text.replace("\u2060", "")  # word joiner
    text = text.replace("\u00ad", "")  # soft hyphen

    # Collapse whitespace left behind by zero-width removals
    text = re.sub(r"[ \t]{2,}", " ", text)  # collapse multiple spaces/tabs
    text = re.sub(r"\n{3,}", "\n\n", text)  # collapse 3+ newlines to 2

    # PDF line-wrap reconstruction (see docstring). We join ONLY mid-sentence
    # line-wraps, not reference-section line breaks. A newline is treated as a
    # line-wrap (joined into a space) only when a lowercase letter precedes it
    # and a lowercase letter follows (allowing optional whitespace on either
    # side — pdfplumber sometimes inserts a leading space or residual zero-width
    # chars at the start of the wrapped line). This preserves:
    #   - Reference-list line breaks (each ref on its own line, typically starts
    #     with an uppercase author name or [n] marker — not "lowercase\nlowercase")
    #   - Heading lines ("References", "Works Cited") that stand alone on a line
    #   - Paragraph breaks (\n\n, untouched)
    # The de-hyphenation step runs first because "represen-\ntation" ends with a
    # lowercase letter before \n and the join would merge it wrongly otherwise.
    text = re.sub(r"-\n(?=\s*[a-z])", "", text)                    # join hyphenated line-breaks
    text = re.sub(r"(?<=[a-z])\s*\n\s*(?=[a-z])", " ", text)        # join mid-sentence wraps only
    # Tidy any double spaces introduced by the join
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text


def extract_from_pdf_pdfplumber(file_path: str) -> str:
    """Extract text from PDF using pdfplumber (good for structured PDFs)."""
    try:
        with pdfplumber.open(file_path) as pdf:
            pages = [page.extract_text() for page in pdf.pages if page.extract_text()]
        return "\n\n".join(pages)
    except Exception as e:
        raise TextExtractionError(f"pdfplumber failed: {e}")


def extract_from_pdf_pymupdf(file_path: str) -> str:
    """Extract text from PDF using PyMuPDF (good fallback)."""
    try:
        doc = fitz.open(file_path)
        pages = [page.get_text() for page in doc]
        doc.close()
        return "\n\n".join(pages)
    except Exception as e:
        raise TextExtractionError(f"PyMuPDF failed: {e}")


def extract_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    try:
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    except Exception as e:
        raise TextExtractionError(f"DOCX extraction failed: {e}")


def extract_text(file_path: str, preferred_backend: str = "pdfplumber") -> str:
    """
    Extract text from a PDF or DOCX file.

    Args:
        file_path: Path to the file.
        preferred_backend: "pdfplumber" or "pymupdf" (PDF only).

    Returns:
        Extracted plain text.

    Raises:
        TextExtractionError if all extraction methods fail.
    """
    path = Path(file_path)
    if not path.exists():
        raise TextExtractionError(f"File not found: {file_path}")

    suffix = path.suffix.lower()

    if suffix == ".pdf":
        # Try preferred backend first, fall back to alternative
        backends = []
        if preferred_backend == "pdfplumber":
            backends = [extract_from_pdf_pdfplumber, extract_from_pdf_pymupdf]
        else:
            backends = [extract_from_pdf_pymupdf, extract_from_pdf_pdfplumber]

        last_error = None
        for backend in backends:
            try:
                text = backend(file_path)
                if text.strip():
                    return _clean_text(text)
            except TextExtractionError as e:
                last_error = e
                logger.warning(f"Backend {backend.__name__} failed: {e}")

        raise TextExtractionError(
            f"All PDF extraction backends failed. Last error: {last_error}"
        )

    elif suffix == ".docx":
        return _clean_text(extract_from_docx(file_path))
    else:
        raise TextExtractionError(f"Unsupported file type: {suffix}")


def extract_text_from_bytes(
    content: bytes, filename: str, preferred_backend: str = "pdfplumber"
) -> str:
    """
    Extract text from file bytes (e.g., from an upload).

    Args:
        content: Raw file bytes.
        filename: Original filename (to determine file type).
        preferred_backend: PDF extraction backend.

    Returns:
        Extracted plain text.
    """
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        # Try pdfplumber first (works with bytes)
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages = [
                    page.extract_text() for page in pdf.pages if page.extract_text()
                ]
            if pages:
                return "\n\n".join(pages)
        except Exception as e:
            logger.warning(f"pdfplumber from bytes failed: {e}")

        # Fallback to PyMuPDF
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            pages = [page.get_text() for page in doc]
            doc.close()
            return "\n\n".join(pages)
        except Exception as e:
            raise TextExtractionError(f"PyMuPDF from bytes failed: {e}")

    elif suffix == ".docx":
        try:
            doc = Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        except Exception as e:
            raise TextExtractionError(f"DOCX from bytes failed: {e}")
    else:
        raise TextExtractionError(f"Unsupported file type: {suffix}")
